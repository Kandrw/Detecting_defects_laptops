import logging
import os
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework import status
from django.http import HttpResponse, JsonResponse

import time


import base64
from django.core.files.base import ContentFile

from detecting_defects_laptops_soft.models import ImageModel


import subprocess
from subprocess import Popen, PIPE
from PIL import Image
from pathlib import Path

from django.core.files.uploadedfile import InMemoryUploadedFile

import io
import PIL
import sys

class ImageUploadView(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        serial_number = request.data.get('serial_number')
        images = request.FILES.getlist('images')
        ARGS_img = ""
        PATH_OUTPUT = "/home/andrey/Документы/Хакатон_13.10.24/Detecting_defects_laptops/mlp/output_txt_files/"
        results = []
        data_img_path = []
        i = 0
        for image in images:
            ImageModel.objects.create(image=image, serial_number=serial_number)
            print("[get] image", image)
            abs_path = "/home/andrey/Документы/Хакатон_13.10.24/Detecting_defects_laptops/Back-hakaton/"
            dir = abs_path + f'media/images/{serial_number}/'
            ARGS_img += str(dir)+str(image) + " "
            print("image",image)
            # print("str(Path(image).stem) ", str(Path(image).stem))
            print("PATH_OUTPUT ", PATH_OUTPUT)
            # tr = PATH_OUTPUT + str(Path(image).stem) + ".txt"
            # print(tr)
            print("str(Path(image.name).stem) ", str(Path(image.name).stem))
            results.append(PATH_OUTPUT + str(Path(image.name).stem) + ".txt")
            # path_img = PATH_OUTPUT + str(image)
            #data_img.append(img_data)
            # with images[i].open() as img_file:
            #     img_content = img_file.read()
            #     encoded_img = base64.b64encode(img_content).decode('utf-8')
            #     img_format = images[i].content_type  # Получаем формат изображения (например, image/png)

            #     # Форматируем изображение в нужный вид
            #     img_data = f"data:{img_format};base64,{encoded_img}"
            # data_img_path.append(img_data)
            print("image -------------------", type(image))
            data_img_path.append(PATH_OUTPUT + str(image.name))
            # results.append(PATH_OUTPUT + "photo_2024-10-09_14-50-32.txt")
            i += 1

        PATH_ML = "../mlp/run_predictions.py"
        PATH_RUN = "/home/andrey/Документы/Хакатон_13.10.24/Detecting_defects_laptops/mlp"
        # InMemoryUploadedFile
        # ARGS = images
        if 1:
            out, err = Popen('python3 ' + PATH_ML + " " + ARGS_img, cwd=PATH_RUN, shell=True, stdout=PIPE).communicate()
            print("ML", out)
        
        result_data_s = []
        errors_str = [
            "Lock",
            "UnknownDefect",
            "MissingScrew",
            "KeysProblems",
            "Chipped",
            "BrokenPixel",
            "Scratch",
            "NoDefect"
        ]

        result_data = {
            "Lock": "",
            "UnknownDefect":"",
            "MissingScrew":"",
            "KeysProblems":"",
            "Chipped":"",
            "BrokenPixel":"",
            "Scratch":"",
            "NoDefect":"",
            "ImgRes":""
        }
        # data_img = []
        # for i in range(len(results)):
        #     with open(results[i], "r") as file:
        #         for line in file:
        #             print(line)
        #             key = (errors_str[int(line[0])])
                    
        #             result_data[key] = result_data[key] + line[1:]
        #             if(result_data[key][-1] == " " or result_data[key][-1] == "\n"):
        #                 result_data[key] = result_data[key][:-1]
            
        #     img_io = io.BytesIO()
        #     # img = Image.open(image)
        #     img_data1 = Image.open(data_img_path[i])
        #     # percent = (basewidth / float(img.size[0]))
        #     # hsize = int(float(img.size[1]) * percent)
        #     # img = img.resize((basewidth, hsize),PIL.Image.ANTIALIAS)
        #     img_data1.save(img_io, format="JPEG")


        #     new_pic= InMemoryUploadedFile(img_io, 
        #         'ImageField',
        #         'profile_pic',
        #         'JPEG',
        #         sys.getsizeof(img_io), None)
            
        #     with new_pic.open() as img_file:
        #         img_content = img_file.read()
        #         encoded_img = base64.b64encode(img_content).decode('utf-8')
        #         img_format = images[i].content_type  # Получаем формат изображения (например, image/png)
        #         # Форматируем изображение в нужный вид
        #         img_data = f"data:{img_format};base64,{encoded_img}"
        #     data_img.append(img_data)
        # result_data["ImgRes"] = data_img
        # print(result_data)

        data_img = []
        for i in range(len(results)):
            with open(results[i], "r") as file:
                for line in file:
                    print(line)
                    key = (errors_str[int(line[0])])
                    result_data[key] += line[1:].strip()  # Используйте strip() для удаления пробелов и новых строк
                    
            img_io = io.BytesIO()  # Перенесите сюда
            img_data1 = Image.open(data_img_path[i])
            img_data1.save(img_io, format="JPEG")
            img_io.seek(0)  # Сбросьте указатель в начале буфера

            new_pic = InMemoryUploadedFile(img_io, 
                'ImageField',
                'profile_pic',
                'JPEG',
                img_io.getbuffer().nbytes, None)  # Измените на img_io.getbuffer().nbytes

            with new_pic.open() as img_file:
                img_content = img_file.read()
                encoded_img = base64.b64encode(img_content).decode('utf-8')
                img_format = images[i].content_type
                img_data = f"data:{img_format};base64,{encoded_img}"
            
            data_img.append(img_data)
        print(f"Number of images processed: {len(data_img)}")  # Выводим количество изображений
        result_data["ImgRes"] = data_img
        # print(result_data)
        return JsonResponse(result_data, status=200)

        # return JsonResponse(result_data, status=200)
class ImageUploadView2(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        serial_number = request.data.get('serial_number')
        images = request.FILES.getlist('images')  # Получаем все изображения
        ARGS_img = ""
        PATH_OUTPUT = "/home/andrey/Документы/Хакатон_13.10.24/Detecting_defects_laptops/mlp/output_txt_files/"
        results = []
        data_img_path = []
        
        # Проходим по каждому изображению
        for image in images:
            ImageModel.objects.create(image=image, serial_number=serial_number)
            abs_path = "/home/andrey/Документы/Хакатон_13.10.24/Detecting_defects_laptops/Back-hakaton/"
            dir = abs_path + f'media/images/{serial_number}/'
            ARGS_img += str(dir) + str(image) + " "
            
            # Добавляем пути для результата
            results.append(PATH_OUTPUT + str(Path(image.name).stem) + ".txt")
            data_img_path.append(dir + str(image))  # Сохраняем полный путь до изображения

        # Запуск модели
        PATH_ML = "../mlp/run_predictions.py"
        PATH_RUN = "/home/andrey/Документы/Хакатон_13.10.24/Detecting_defects_laptops/mlp"
        
        out, err = Popen('python3 ' + PATH_ML + " " + ARGS_img, cwd=PATH_RUN, shell=True, stdout=PIPE).communicate()
        print("ML Output:", out)
        
        result_data_s = []
        errors_str = [
            "Lock", "UnknownDefect", "MissingScrew", "KeysProblems",
            "Chipped", "BrokenPixel", "Scratch", "NoDefect"
        ]
        
        result_data = {
            "Lock": "", "UnknownDefect": "", "MissingScrew": "", "KeysProblems": "",
            "Chipped": "", "BrokenPixel": "", "Scratch": "", "NoDefect": "", "ImgRes": []
        }
        
        # Подготовка изображений для ответа
        data_img = []
        for i, result_path in enumerate(results):
            with open(result_path, "r") as file:
                for line in file:
                    key = errors_str[int(line[0])]
                    result_data[key] = result_data[key] + line[1:].strip()
            
            img_io = io.BytesIO()
            img_data = Image.open(data_img_path[i])  # Открываем изображение
            img_data.save(img_io, format="JPEG")
            
            # Получаем контент изображения и кодируем его в base64
            img_io.seek(0)
            img_content = img_io.read()
            encoded_img = base64.b64encode(img_content).decode('utf-8')
            img_format = "image/jpeg"  # Если все изображения в формате JPEG

            # Добавляем изображение в результат
            img_data_base64 = f"data:{img_format};base64,{encoded_img}"
            data_img.append(img_data_base64)
        
        result_data["ImgRes"] = data_img  # Все изображения в base64 формате
        print("=========", len(data_img))
        return JsonResponse(result_data, status=200)


class ImageUploadView3(APIView):
    parser_classes = (MultiPartParser, FormParser)

    def post(self, request, *args, **kwargs):
        serial_number = request.data.get('serial_number')
        images = request.FILES.getlist('images')  # Получаем все изображения
        ARGS_img = ""
        PATH_OUTPUT = "/home/andrey/Документы/Хакатон_13.10.24/Detecting_defects_laptops/mlp/output_txt_files/"
        results = []
        data_img_path = []
        
        # Проходим по каждому изображению
        for image in images:
            ImageModel.objects.create(image=image, serial_number=serial_number)
            abs_path = "/home/andrey/Документы/Хакатон_13.10.24/Detecting_defects_laptops/Back-hakaton/"
            dir = abs_path + f'media/images/{serial_number}/'
            ARGS_img += str(dir) + str(image) + " "
            
            # Добавляем пути для результата
            results.append(PATH_OUTPUT + str(Path(image.name).stem) + ".txt")
            data_img_path.append(dir + str(image))  # Сохраняем полный путь до изображения

        # Запуск модели
        PATH_ML = "../mlp/run_predictions.py"
        PATH_RUN = "/home/andrey/Документы/Хакатон_13.10.24/Detecting_defects_laptops/mlp"
        
        out, err = Popen('python3 ' + PATH_ML + " " + ARGS_img, cwd=PATH_RUN, shell=True, stdout=PIPE).communicate()
        print("ML Output:", out)
        
        result_data_s = []
        errors_str = [
            "Lock", "UnknownDefect", "MissingScrew", "KeysProblems",
            "Chipped", "BrokenPixel", "Scratch", "NoDefect"
        ]
        
        result_data = {
            "Lock": "", "UnknownDefect": "", "MissingScrew": "", "KeysProblems": "",
            "Chipped": "", "BrokenPixel": "", "Scratch": "", "NoDefect": "", "ImgRes": []
        }
        
        # Подготовка изображений для ответа
        data_img = []
        for i, result_path in enumerate(results):
            with open(result_path, "r") as file:
                for line in file:
                    key = errors_str[int(line[0])]
                    result_data[key] = result_data[key] + line[1:].strip()
            
            img_io = io.BytesIO()
            img_data = Image.open(data_img_path[i])  # Открываем изображение
            img_data.save(img_io, format="JPEG")
            
            # Получаем контент изображения и кодируем его в base64
            img_io.seek(0)
            img_content = img_io.read()
            encoded_img = base64.b64encode(img_content).decode('utf-8')
            img_format = "image/jpeg"  # Если все изображения в формате JPEG

            # Добавляем изображение в результат
            img_data_base64 = f"data:{img_format};base64,{encoded_img}"
            data_img.append(img_data_base64)
        
        result_data["ImgRes"] = data_img  # Все изображения в base64 формате
        
        return JsonResponse(result_data, status=200)

from django.http import JsonResponse
from rest_framework.decorators import api_view
from django.core.files.storage import default_storage
from docx import Document

@api_view(['POST'])
def submit_report(request):
    serial_number = request.data.get('serialNumber')
    defects = request.data.get('defects')

    if not defects:
        return JsonResponse({'error': 'Defects data is missing or empty'}, status=400)

    doc = Document()
    doc.add_heading('Отчет по дефектам', level=1)
    doc.add_paragraph(f'Серийный номер: {serial_number}')

    # for key, value in defects.items():
    #     doc.add_paragraph(f'{key}: {value}')
    for key, value in list(defects.items())[:-1]:
        doc.add_paragraph(f'{key}: {value}')
    file_path = f'reports/report_{serial_number}.docx'
    doc.save(file_path)

    with open(file_path, 'rb') as f:
        response = JsonResponse({'status': 'success'})
        response['Content-Disposition'] = f'attachment; filename="report_{serial_number}.docx"'
        response.write(f.read())
        return response

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

@api_view(['POST'])
def submit_report_pdf(request):
    serial_number = request.data.get('serialNumber')
    defects = request.data.get('defects')

    if not defects:
        return JsonResponse({'error': 'Defects data is missing or empty'}, status=400)

    file_path = f'reports/report_{serial_number}.pdf'
    c = canvas.Canvas(file_path, pagesize=letter)
    c.drawString(100, 750, f'Report about Laptop defects')
    c.drawString(100, 735, f'Serial Number: {serial_number}')

    y = 700
    # for key, value in defects.items():
    #     c.drawString(100, y, f'{key}: {value}')
    #     y -= 20
    for key, value in list(defects.items())[:-1]:
        c.drawString(100, y, f'{key}: {value}')
        y -= 20
    c.save()

    with open(file_path, 'rb') as f:
        response = JsonResponse({'status': 'success'})
        response['Content-Disposition'] = f'attachment; filename="report_{serial_number}.pdf"'
        response.write(f.read())
        return response
